import streamlit as st
from io import BytesIO
import re
from typing import Dict, List, Any, Optional
import docx
import PyPDF2
import pdfplumber
import logging

class DocumentParser:
    """
    Handles document ingestion and basic structure detection for PDF, DOCX, and TXT files.
    """
    
    def __init__(self):
        self.section_patterns = [
            r'(?i)\b(?:scope\s+of\s+work|scope|project\s+scope)\b',
            r'(?i)\b(?:deliverables?|outcomes?|outputs?)\b',
            r'(?i)\b(?:timeline|schedule|milestones?|deadlines?)\b',
            r'(?i)\b(?:acceptance\s+criteria|acceptance|testing)\b',
            r'(?i)\b(?:assumptions?|dependencies|constraints?)\b',
            r'(?i)\b(?:pricing|payment\s+terms|cost|budget)\b',
            r'(?i)\b(?:change\s+control|changes?|modifications?)\b',
            r'(?i)\b(?:objectives?|goals?|purpose)\b'
        ]
    
    def parse_document(self, uploaded_file) -> Dict[str, Any]:
        """
        Main method to parse uploaded document and extract structure.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dictionary containing parsed document data
        """
        try:
            # Get file extension
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            # Extract text based on file type
            if file_extension == 'pdf':
                text_content = self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                text_content = self._extract_from_docx(uploaded_file)
            elif file_extension == 'txt':
                text_content = self._extract_from_txt(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Process the extracted text
            document_data = self._process_text(text_content, uploaded_file.name)
            
            return document_data
            
        except Exception as e:
            logging.error(f"Error parsing document: {str(e)}")
            raise Exception(f"Failed to parse document: {str(e)}")
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file using pdfplumber with PyPDF2 fallback."""
        text_content = ""
        
        # First try with pdfplumber (more reliable)
        try:
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                return text_content
        except Exception as e:
            logging.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
            
            if text_content.strip():
                return text_content
                
        except Exception as e:
            logging.warning(f"PyPDF2 extraction also failed: {str(e)}")
        
        # If both methods fail
        if not text_content.strip():
            raise ValueError("No text could be extracted from PDF. The PDF might be image-based or encrypted. Please try a different PDF or convert it to text format.")
            
        return text_content
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(uploaded_file)
            text_content = ""
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            if not text_content.strip():
                raise ValueError("No text could be extracted from DOCX")
                
            return text_content
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file."""
        try:
            # Read as string
            text_content = str(uploaded_file.read(), "utf-8")
            
            if not text_content.strip():
                raise ValueError("Text file is empty")
                
            return text_content
            
        except Exception as e:
            raise Exception(f"TXT extraction failed: {str(e)}")
    
    def _process_text(self, text_content: str, filename: str) -> Dict[str, Any]:
        """
        Process extracted text to identify structure and sections.
        
        Args:
            text_content: Raw text from document
            filename: Original filename
            
        Returns:
            Structured document data
        """
        # Split into paragraphs
        paragraphs = [p.strip() for p in text_content.split('\n') if p.strip()]
        
        # Detect sections
        sections = self._detect_sections(text_content)
        
        # Detect formatting elements
        formatting_info = self._analyze_formatting(text_content)
        
        # Count words and characters
        word_count = len(text_content.split())
        char_count = len(text_content)
        
        document_data = {
            'filename': filename,
            'raw_text': text_content,
            'paragraphs': paragraphs,
            'sections': sections,
            'formatting': formatting_info,
            'word_count': word_count,
            'character_count': char_count,
            'paragraph_count': len(paragraphs)
        }
        
        return document_data
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Detect document sections based on common SoW patterns.
        
        Args:
            text: Document text
            
        Returns:
            Dictionary mapping section names to their content
        """
        sections = {}
        section_names = [
            'scope', 'deliverables', 'timeline', 'acceptance_criteria',
            'assumptions', 'pricing', 'change_control', 'objectives'
        ]
        
        # Split text into potential sections based on headings
        lines = text.split('\n')
        current_section = 'general'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line looks like a heading
            is_heading = self._is_potential_heading(line)
            
            if is_heading:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Identify section type
                section_type = self._identify_section_type(line)
                current_section = section_type
                current_content = []
            else:
                current_content.append(line)
        
        # Save final section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _is_potential_heading(self, line: str) -> bool:
        """Check if a line looks like a section heading."""
        # Check for common heading patterns
        heading_patterns = [
            r'^\d+\.?\s+',  # Numbered headings (1. or 1)
            r'^[A-Z][A-Z\s]+:?$',  # ALL CAPS headings
            r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*:?$',  # Title Case headings
            r'^\w+\s*:\s*$'  # Word followed by colon
        ]
        
        return any(re.match(pattern, line) for pattern in heading_patterns)
    
    def _identify_section_type(self, heading: str) -> str:
        """Identify the type of section based on heading text."""
        heading_lower = heading.lower()
        
        section_mapping = {
            'scope': ['scope', 'project scope', 'scope of work'],
            'objectives': ['objective', 'objectives', 'goal', 'goals', 'purpose'],
            'deliverables': ['deliverable', 'deliverables', 'outcome', 'outcomes', 'output', 'outputs'],
            'timeline': ['timeline', 'schedule', 'milestone', 'milestones', 'deadline', 'deadlines'],
            'acceptance_criteria': ['acceptance', 'acceptance criteria', 'testing', 'acceptance test'],
            'assumptions': ['assumption', 'assumptions', 'dependencies', 'dependency', 'constraint', 'constraints'],
            'pricing': ['pricing', 'payment', 'payment terms', 'cost', 'budget', 'fee', 'fees'],
            'change_control': ['change', 'changes', 'change control', 'modification', 'modifications']
        }
        
        for section_type, keywords in section_mapping.items():
            for keyword in keywords:
                if keyword in heading_lower:
                    return section_type
        
        return 'general'
    
    def _analyze_formatting(self, text: str) -> Dict[str, Any]:
        """
        Analyze formatting elements in the text.
        
        Args:
            text: Document text
            
        Returns:
            Dictionary with formatting analysis
        """
        # Count bullet points
        bullet_patterns = [r'•', r'-\s+', r'\*\s+', r'◦']
        bullet_count = sum(len(re.findall(pattern, text)) for pattern in bullet_patterns)
        
        # Count numbered lists
        numbered_list_count = len(re.findall(r'^\d+\.', text, re.MULTILINE))
        
        # Check for tables (basic detection)
        table_indicators = ['|', '\t\t', '    ']  # Common table separators
        potential_tables = any(indicator in text for indicator in table_indicators)
        
        # Check for consistent capitalization
        words = re.findall(r'\b[A-Za-z]+\b', text)
        inconsistent_caps = len(set(word.lower() for word in words if len(word) > 3)) != len(set(words))
        
        formatting_info = {
            'has_bullet_points': bullet_count > 0,
            'bullet_count': bullet_count,
            'has_numbered_lists': numbered_list_count > 0,
            'numbered_list_count': numbered_list_count,
            'has_potential_tables': potential_tables,
            'inconsistent_capitalization': inconsistent_caps,
            'line_count': len(text.split('\n'))
        }
        
        return formatting_info
